"""
Generate a professional requirements document (.docx) from structured JSON.
Uses Emergent branding but IEEE 830 / ISO 29148 inspired structure.

Each requirement is a full section with: Description, Rationale, Source/Citation,
Priority, Acceptance Criteria, Dependencies, and Validation Status.
"""

import json
import sys
import os
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# Emergent branding
BRAND_PRIMARY = "EE3342"  # coral/red
BRAND_NAVY = "333F4F"
BRAND_LIGHT = "FDF2F2"
BRAND_ACCENT_LIGHT = "F8F8F8"
CALLOUT_BG = "FFF8E1"
CALLOUT_BORDER = "F9A825"


def rgb(hex_color):
    return RGBColor.from_string(hex_color)


def set_cell_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._element.get_or_add_tcPr().append(shading)


def set_cell_margins(cell, top=0, bottom=0, start=0, end=0):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:start w:w="{start}" w:type="dxa"/>'
        f'  <w:end w:w="{end}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)


def set_cell_border_bottom(cell, color, sz=8):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)


def set_cell_border_top(cell, color, sz=8):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)


def add_run(paragraph, text, font_name="Calibri", font_size=11, bold=False, color=None, italic=False):
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = rgb(color)
    return run


def format_bullet(paragraph):
    """Apply correct bullet paragraph formatting:
    Left indent 0.25", Hanging 0.25", Before/After 6pt, Line spacing At Least 12pt."""
    fmt = paragraph.paragraph_format
    fmt.left_indent = Inches(0.25)
    fmt.first_line_indent = Inches(-0.25)
    fmt.space_before = Pt(6)
    fmt.space_after = Pt(6)
    fmt.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    fmt.line_spacing = Pt(12)


def add_bullet(doc, text, font_size=11):
    """Add a properly formatted bullet paragraph."""
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    add_run(p, text, font_size=font_size)
    format_bullet(p)
    return p


def add_styled_table(doc, headers, rows, col_widths=None):
    """Create a slim professional table."""
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = True

    # Header row
    for ci, header in enumerate(headers):
        cell = tbl.rows[0].cells[ci]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        run.font.color.rgb = rgb(BRAND_PRIMARY)
        set_cell_border_bottom(cell, BRAND_PRIMARY)
        set_cell_margins(cell, top=40, bottom=40, start=80, end=80)

    # Data rows
    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri + 1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(11)
            run.font.name = "Calibri"
            set_cell_margins(cell, top=30, bottom=30, start=80, end=80)
            if ri % 2 == 0:
                set_cell_shading(cell, BRAND_ACCENT_LIGHT)

    # Column widths
    if col_widths:
        for row in tbl.rows:
            for ci, w in enumerate(col_widths):
                if ci < len(row.cells):
                    row.cells[ci].width = Inches(w)

    # Repeat header row
    tbl.rows[0].height_rule = 1
    tr = tbl.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = parse_xml(f'<w:tblHeader {nsdecls("w")}/>')
    trPr.append(tblHeader)

    return tbl


def add_requirement_block(doc, req):
    """Add a full requirement section with all detail fields."""
    # Requirement ID and title as heading
    h = doc.add_heading(level=3)
    add_run(h, f"{req['id']}: ", font_size=12, bold=True, color=BRAND_PRIMARY)
    add_run(h, req["title"], font_size=12, bold=True, color=BRAND_NAVY)

    # Detail table (2 columns: label, value)
    details = []

    # Description
    p = doc.add_paragraph()
    add_run(p, "Description: ", font_size=11, bold=True, color=BRAND_NAVY)
    add_run(p, req["description"], font_size=11)

    # Rationale
    if req.get("rationale"):
        p = doc.add_paragraph()
        add_run(p, "Rationale: ", font_size=11, bold=True, color=BRAND_NAVY)
        add_run(p, req["rationale"], font_size=11)

    # Source / Citation
    if req.get("source"):
        p = doc.add_paragraph()
        add_run(p, "Source: ", font_size=11, bold=True, color=BRAND_NAVY)
        add_run(p, req["source"], font_size=11, italic=True)

    # Priority
    priority_colors = {"Must": "C62828", "Should": "E65100", "Could": "2E7D32", "Won't": "616161"}
    if req.get("priority"):
        p = doc.add_paragraph()
        add_run(p, "Priority: ", font_size=11, bold=True, color=BRAND_NAVY)
        pc = priority_colors.get(req["priority"], BRAND_NAVY)
        add_run(p, req["priority"], font_size=11, bold=True, color=pc)

    # Acceptance Criteria
    if req.get("acceptance_criteria"):
        p = doc.add_paragraph()
        add_run(p, "Acceptance Criteria:", font_size=11, bold=True, color=BRAND_NAVY)
        for ac in req["acceptance_criteria"]:
            add_bullet(doc, ac)

    # Dependencies
    if req.get("dependencies"):
        p = doc.add_paragraph()
        add_run(p, "Dependencies: ", font_size=11, bold=True, color=BRAND_NAVY)
        add_run(p, req["dependencies"], font_size=11)

    # Validation Status
    if req.get("validated"):
        p = doc.add_paragraph()
        add_run(p, "Validation Status: ", font_size=11, bold=True, color=BRAND_NAVY)
        v = req["validated"]
        vc = "2E7D32" if v == "Confirmed" else ("E65100" if v == "Assumption" else "616161")
        add_run(p, v, font_size=11, bold=True, color=vc)

    # Thin separator
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)


def generate_requirements_doc(content):
    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # ---- TITLE PAGE ----
    for _ in range(3):
        doc.add_paragraph()

    # Project name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, content["project_name"], font_size=28, bold=True, color=BRAND_NAVY)

    # Thin accent line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    cell.text = ""
    set_cell_shading(cell, BRAND_PRIMARY)
    cell.width = Inches(4)
    row = tbl.rows[0]
    row.height = Inches(0.04)

    # Document title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, content["document_title"], font_size=20, color=BRAND_NAVY)

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "Software Requirements Specification", font_size=14, color="666666", italic=True)

    doc.add_paragraph()

    # Document control table
    control_data = [
        ["Version", content.get("version", "1.0")],
        ["Date", content.get("date", datetime.now().strftime("%B %d, %Y"))],
        ["Prepared By", content.get("prepared_by", "Emergent Software")],
        ["Status", content.get("status", "Draft")],
    ]
    for label, value in control_data:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, f"{label}: ", font_size=11, bold=True, color=BRAND_NAVY)
        add_run(p, value, font_size=11)

    # Page break
    doc.add_page_break()

    # ---- REVISION HISTORY ----
    doc.add_heading("Revision History", level=1).runs[0].font.color.rgb = rgb(BRAND_PRIMARY)

    if content.get("revision_history"):
        add_styled_table(doc,
            ["Version", "Date", "Description", "Author"],
            content["revision_history"],
            col_widths=[0.7, 1.2, 4.0, 1.5]
        )
    doc.add_paragraph()

    # ---- TABLE OF CONTENTS placeholder ----
    doc.add_heading("Table of Contents", level=1).runs[0].font.color.rgb = rgb(BRAND_PRIMARY)
    p = doc.add_paragraph()
    add_run(p, "[Table of Contents \u2014 update field after opening in Word]", font_size=11, italic=True, color="999999")
    doc.add_page_break()

    # ---- 1. INTRODUCTION ----
    doc.add_heading("1. Introduction", level=1).runs[0].font.color.rgb = rgb(BRAND_PRIMARY)

    doc.add_heading("1.1 Purpose", level=2).runs[0].font.color.rgb = rgb(BRAND_NAVY)
    p = doc.add_paragraph()
    add_run(p, content.get("purpose", "This document defines the software requirements for the described feature set. It is intended to serve as the authoritative reference for what the system must do, why, and under what constraints."), font_size=11)

    doc.add_heading("1.2 Scope", level=2).runs[0].font.color.rgb = rgb(BRAND_NAVY)
    p = doc.add_paragraph()
    add_run(p, content.get("scope", ""), font_size=11)

    if content.get("definitions"):
        doc.add_heading("1.3 Definitions & Acronyms", level=2).runs[0].font.color.rgb = rgb(BRAND_NAVY)
        add_styled_table(doc,
            ["Term", "Definition"],
            content["definitions"],
            col_widths=[2.0, 5.0]
        )

    doc.add_heading("1.4 References", level=2).runs[0].font.color.rgb = rgb(BRAND_NAVY)
    if content.get("sources"):
        add_styled_table(doc,
            ["ID", "Type", "Description", "Author", "Date"],
            [[s["id"], s["type"], s["description"], s["who"], s["date"]] for s in content["sources"]],
            col_widths=[0.4, 0.8, 3.2, 1.5, 1.0]
        )

    doc.add_page_break()

    # ---- 2. OVERALL DESCRIPTION ----
    doc.add_heading("2. Overall Description", level=1).runs[0].font.color.rgb = rgb(BRAND_PRIMARY)

    doc.add_heading("2.1 Product Perspective", level=2).runs[0].font.color.rgb = rgb(BRAND_NAVY)
    p = doc.add_paragraph()
    add_run(p, content.get("product_perspective", ""), font_size=11)

    if content.get("user_characteristics"):
        doc.add_heading("2.2 User Characteristics", level=2).runs[0].font.color.rgb = rgb(BRAND_NAVY)
        for uc in content["user_characteristics"]:
            add_bullet(doc, uc)

    doc.add_heading("2.3 Assumptions", level=2).runs[0].font.color.rgb = rgb(BRAND_NAVY)
    p = doc.add_paragraph()
    add_run(p, "The following assumptions have been made during requirements analysis. Each introduces risk if incorrect and should be validated with stakeholders before implementation begins.", font_size=11, italic=True)
    doc.add_paragraph()

    if content.get("assumptions"):
        add_styled_table(doc,
            ["ID", "Assumption", "Risk if Wrong", "Source"],
            [[a["id"], a["assumption"], a["risk"], a["source"]] for a in content["assumptions"]],
            col_widths=[0.4, 3.2, 2.2, 1.5]
        )

    if content.get("constraints"):
        doc.add_heading("2.4 Constraints & Dependencies", level=2).runs[0].font.color.rgb = rgb(BRAND_NAVY)
        for c in content["constraints"]:
            add_bullet(doc, c)

    doc.add_page_break()

    # ---- 3. SPECIFIC REQUIREMENTS ----
    doc.add_heading("3. Specific Requirements", level=1).runs[0].font.color.rgb = rgb(BRAND_PRIMARY)

    # Group requirements by category
    for category in content.get("requirement_categories", []):
        doc.add_heading(f"3.{category['number']} {category['name']}", level=2).runs[0].font.color.rgb = rgb(BRAND_NAVY)

        if category.get("overview"):
            p = doc.add_paragraph()
            add_run(p, category["overview"], font_size=11)
            doc.add_paragraph()

        for req in category.get("requirements", []):
            add_requirement_block(doc, req)

    doc.add_page_break()

    # ---- 4. NON-FUNCTIONAL REQUIREMENTS ----
    doc.add_heading("4. Non-Functional Requirements", level=1).runs[0].font.color.rgb = rgb(BRAND_PRIMARY)

    for req in content.get("nfrs", []):
        add_requirement_block(doc, req)

    # ---- 5. OPEN QUESTIONS ----
    doc.add_heading("5. Open Questions", level=1).runs[0].font.color.rgb = rgb(BRAND_PRIMARY)

    if content.get("resolved_questions"):
        doc.add_heading("5.1 Resolved", level=2).runs[0].font.color.rgb = rgb(BRAND_NAVY)
        add_styled_table(doc,
            ["#", "Question", "Answer", "Answered By", "Impact"],
            content["resolved_questions"],
            col_widths=[0.3, 2.5, 2.5, 1.0, 1.0]
        )
        doc.add_paragraph()

    if content.get("unresolved_questions"):
        doc.add_heading("5.2 Unresolved", level=2).runs[0].font.color.rgb = rgb(BRAND_NAVY)

        # Callout box for urgency
        callout_tbl = doc.add_table(rows=1, cols=1)
        callout_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = callout_tbl.rows[0].cells[0]
        cell.text = ""
        set_cell_shading(cell, CALLOUT_BG)
        set_cell_margins(cell, top=100, bottom=100, start=150, end=150)
        p = cell.paragraphs[0]
        add_run(p, "\u26a0 ", font_size=11)
        add_run(p, f"These {len(content['unresolved_questions'])} questions must be resolved before implementation can begin.", font_size=11, bold=True)
        doc.add_paragraph()

        add_styled_table(doc,
            ["#", "Question", "Category", "Needed From", "Why It Matters", "Blocking?"],
            content["unresolved_questions"],
            col_widths=[0.3, 2.0, 0.8, 1.0, 2.0, 0.6]
        )

    doc.add_page_break()

    # ---- 6. TRACEABILITY MATRIX ----
    doc.add_heading("6. Traceability Matrix", level=1).runs[0].font.color.rgb = rgb(BRAND_PRIMARY)
    p = doc.add_paragraph()
    add_run(p, "Maps each requirement back to its source and forward to suggested implementation stories.", font_size=11, italic=True)
    doc.add_paragraph()

    if content.get("traceability"):
        add_styled_table(doc,
            ["Req ID", "Source", "Priority", "Validated", "Suggested Story"],
            content["traceability"],
            col_widths=[0.6, 1.5, 0.7, 0.8, 3.0]
        )

    doc.add_page_break()

    # ---- 7. SUGGESTED STORY DECOMPOSITION ----
    doc.add_heading("7. Suggested Story Decomposition", level=1).runs[0].font.color.rgb = rgb(BRAND_PRIMARY)
    p = doc.add_paragraph()
    add_run(p, "Based on the requirements above, the following stories are suggested. Point estimates use the project's Fibonacci scale. Each story is a cohesive, independently deliverable unit of work.", font_size=11, italic=True)
    doc.add_paragraph()

    for story in content.get("stories", []):
        h = doc.add_heading(level=3)
        add_run(h, f"Story {story['number']}: ", font_size=12, bold=True, color=BRAND_PRIMARY)
        add_run(h, story["title"], font_size=12, bold=True, color=BRAND_NAVY)
        add_run(h, f"  (~{story['points']} pts)", font_size=11, color="666666")

        # Description
        if story.get("description"):
            p = doc.add_paragraph()
            add_run(p, story["description"], font_size=11)

        p = doc.add_paragraph()
        add_run(p, "Covers: ", font_size=11, bold=True, color=BRAND_NAVY)
        add_run(p, story.get("covers", ""), font_size=11)

        if story.get("acceptance_criteria"):
            for ac in story["acceptance_criteria"]:
                add_bullet(doc, ac)
        doc.add_paragraph()

    # Summary table
    if content.get("stories"):
        doc.add_heading("Estimate Summary", level=2).runs[0].font.color.rgb = rgb(BRAND_NAVY)
        story_rows = [[s["number"], s["title"], str(s["points"])] for s in content["stories"]]
        total_pts = sum(s["points"] for s in content["stories"])
        story_rows.append(["", "Total", str(total_pts)])
        add_styled_table(doc,
            ["#", "Story", "Points"],
            story_rows,
            col_widths=[0.4, 5.5, 0.8]
        )

    doc.add_page_break()

    # ---- 8. SIGN-OFF ----
    doc.add_heading("8. Approval & Sign-Off", level=1).runs[0].font.color.rgb = rgb(BRAND_PRIMARY)
    p = doc.add_paragraph()
    add_run(p, "By signing below, the undersigned acknowledges that they have reviewed this requirements document and agree that it accurately represents the intended scope of work.", font_size=11)
    doc.add_paragraph()
    doc.add_paragraph()

    signoff_table = doc.add_table(rows=4, cols=4)
    signoff_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Role", "Name", "Signature", "Date"]
    roles = ["Product Owner", "Project Manager", "Technical Lead"]
    for ci, h in enumerate(headers):
        cell = signoff_table.rows[0].cells[ci]
        cell.text = ""
        p = cell.paragraphs[0]
        add_run(p, h, font_size=11, bold=True, color=BRAND_NAVY)
        set_cell_border_bottom(cell, BRAND_PRIMARY)
    for ri, role in enumerate(roles):
        cell = signoff_table.rows[ri + 1].cells[0]
        cell.text = ""
        p = cell.paragraphs[0]
        add_run(p, role, font_size=11)
        for ci in range(1, 4):
            cell = signoff_table.rows[ri + 1].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            add_run(p, "________________________", font_size=11, color="CCCCCC")

    # ---- FOOTER ----
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.clear()
        run = p.add_run(f"{content['project_name']} \u2022 Emergent Software \u2022 Confidential")
        run.font.size = Pt(8)
        run.font.color.rgb = rgb("999999")
        run.font.name = "Calibri"

        # Page numbers
        p2 = footer.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run()
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run2._element.append(fldChar1)
        run3 = p2.add_run()
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run3._element.append(instrText)
        run4 = p2.add_run()
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run4._element.append(fldChar2)
        for r in [run2, run3, run4]:
            r.font.size = Pt(8)
            r.font.color.rgb = rgb("999999")

    # Save
    output_path = content["output_path"]
    doc.save(output_path)
    print(f"Requirements document saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python generate_requirements_doc.py content.json")
        sys.exit(1)
    with open(sys.argv[1], encoding="utf-8") as f:
        data = json.load(f)
    generate_requirements_doc(data)
