"""
Generate a professional software effort estimate document (.docx).

This is a bid/proposal-style document — distinct from requirements specs
and status reports. The structure follows industry estimation best practices
(PERT, WBS, phase-level costing, risk-adjusted totals).

Usage:
    python generate_estimate_doc.py estimate_content.json
"""

import json
import sys
import os
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ── Emergent branding ──────────────────────────────────────────────────
BRAND_PRIMARY = "EE3342"
BRAND_NAVY = "333F4F"
BRAND_LIGHT = "FDF2F2"
BRAND_ALT_ROW = "F8F8F8"
CALLOUT_BG = "FFF8E1"
CALLOUT_BORDER = "F9A825"
FONT = "Calibri"
BODY_SIZE = 11
TABLE_SIZE = 11
HEADER_SIZE = 11


# ── Helpers ────────────────────────────────────────────────────────────

def rgb(hex_color):
    return RGBColor.from_string(hex_color)


def set_shading(cell, color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._element.get_or_add_tcPr().append(shading)


def set_margins(cell, top=40, bottom=40, start=80, end=80):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:start w:w="{start}" w:type="dxa"/>'
        f'  <w:end w:w="{end}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)


def set_border_bottom(cell, color, sz=8):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    b = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(b)


def set_border_top(cell, color, sz=8):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    b = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(b)


def run(paragraph, text, size=BODY_SIZE, bold=False, color=None, italic=False, font=FONT):
    r = paragraph.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = rgb(color)
    return r


def bullet(doc, text, size=BODY_SIZE):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    run(p, text, size=size)
    fmt = p.paragraph_format
    fmt.left_indent = Inches(0.25)
    fmt.first_line_indent = Inches(-0.25)
    fmt.space_before = Pt(6)
    fmt.space_after = Pt(6)
    fmt.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    fmt.line_spacing = Pt(12)
    return p


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = rgb(BRAND_PRIMARY if level == 1 else BRAND_NAVY)
    return h


def para(doc, text, size=BODY_SIZE, bold=False, color=None, italic=False):
    p = doc.add_paragraph()
    run(p, text, size=size, bold=bold, color=color, italic=italic)
    return p


def repeat_header(row):
    """Mark a table row to repeat as header on each page."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = parse_xml(f'<w:tblHeader {nsdecls("w")}/>')
    trPr.append(tblHeader)


def slim_table(doc, headers, rows, col_widths=None, right_align_cols=None):
    """Create a slim professional table with thin header border, alternating rows."""
    right_align_cols = right_align_cols or []
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = True

    # Header
    for ci, h in enumerate(headers):
        cell = tbl.rows[0].cells[ci]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if ci in right_align_cols else WD_ALIGN_PARAGRAPH.LEFT
        run(p, h, size=TABLE_SIZE, bold=True, color=BRAND_PRIMARY)
        set_border_bottom(cell, BRAND_PRIMARY)
        set_margins(cell)

    repeat_header(tbl.rows[0])

    # Data rows
    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri + 1]
        is_total = row_data.get("_total", False) if isinstance(row_data, dict) else False
        vals = row_data.get("_vals", row_data) if isinstance(row_data, dict) else row_data

        for ci, val in enumerate(vals):
            cell = row.cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if ci in right_align_cols else WD_ALIGN_PARAGRAPH.LEFT
            run(p, str(val), size=TABLE_SIZE, bold=is_total)
            set_margins(cell)

            if is_total:
                set_border_top(cell, BRAND_PRIMARY, sz=6)
            elif ri % 2 == 0:
                set_shading(cell, BRAND_ALT_ROW)

    if col_widths:
        for row in tbl.rows:
            for ci, w in enumerate(col_widths):
                if ci < len(row.cells):
                    row.cells[ci].width = Inches(w)

    return tbl


def callout(doc, text):
    """Yellow callout box."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    cell.text = ""
    set_shading(cell, CALLOUT_BG)
    set_margins(cell, top=100, bottom=100, start=150, end=150)
    p = cell.paragraphs[0]
    run(p, text, size=BODY_SIZE)


def accent_line(doc):
    """Thin brand-colored accent line."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    cell.text = ""
    set_shading(cell, BRAND_PRIMARY)
    cell.width = Inches(5)
    tbl.rows[0].height = Inches(0.04)


# ── Main generator ─────────────────────────────────────────────────────

def generate_estimate(content):
    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # ════════════════════════════════════════════════════════════════════
    # TITLE PAGE
    # ════════════════════════════════════════════════════════════════════
    for _ in range(3):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, content["project_name"], size=28, bold=True, color=BRAND_NAVY)

    accent_line(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, "Effort Estimate", size=20, color=BRAND_NAVY)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(p, content.get("feature_name", ""), size=16, color=BRAND_PRIMARY)

    doc.add_paragraph()

    # Document control
    control = [
        ("Version", content.get("version", "1.0")),
        ("Date", content.get("date", datetime.now().strftime("%B %d, %Y"))),
        ("Prepared By", content.get("prepared_by", "Emergent Software")),
        ("Delivery Model", content.get("delivery_model", "Agentic (AI-Assisted)")),
        ("Status", content.get("status", "Draft")),
        ("Confidence", content.get("overall_confidence", "Medium")),
    ]
    for label, value in control:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run(p, f"{label}: ", size=BODY_SIZE, bold=True, color=BRAND_NAVY)
        run(p, value, size=BODY_SIZE)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # ESTIMATE AT A GLANCE
    # ════════════════════════════════════════════════════════════════════
    heading(doc, "Estimate at a Glance")

    glance = content.get("at_a_glance", {})
    slim_table(doc,
        ["Metric", "Value"],
        [
            [k, v] for k, v in glance.items()
        ],
        col_widths=[3.5, 3.5]
    )

    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════
    # EXECUTIVE SUMMARY
    # ════════════════════════════════════════════════════════════════════
    heading(doc, "Executive Summary")

    for paragraph_text in content.get("executive_summary", []):
        para(doc, paragraph_text)

    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════
    # METHODOLOGY
    # ════════════════════════════════════════════════════════════════════
    heading(doc, "Estimation Methodology")

    for section in content.get("methodology", []):
        heading(doc, section["title"], level=2)
        para(doc, section["text"])
        if section.get("table"):
            t = section["table"]
            slim_table(doc, t["headers"], t["rows"],
                       col_widths=t.get("col_widths"),
                       right_align_cols=t.get("right_align_cols"))
            doc.add_paragraph()

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # DETAILED ESTIMATES BY STORY
    # ════════════════════════════════════════════════════════════════════
    heading(doc, "Detailed Estimates")

    for feature in content.get("features", []):
        heading(doc, feature["name"], level=2)
        para(doc, feature["overview"])
        doc.add_paragraph()

        for story in feature.get("stories", []):
            # Story heading
            h = doc.add_heading(level=3)
            run(h, f"Story {story['number']}: ", size=12, bold=True, color=BRAND_PRIMARY)
            run(h, story["title"], size=12, bold=True, color=BRAND_NAVY)

            # Hours table
            slim_table(doc,
                ["", "Optimistic", "Most Likely", "Pessimistic", "PERT Expected"],
                [
                    ["Traditional", story["trad_o"], story["trad_m"], story["trad_p"], story["trad_pert"]],
                    ["Agentic", story["agnt_o"], story["agnt_m"], story["agnt_p"], story["agnt_pert"]],
                ],
                col_widths=[1.2, 1.2, 1.2, 1.2, 1.2],
                right_align_cols=[1, 2, 3, 4]
            )

            doc.add_paragraph()

            # What this involves
            p = doc.add_paragraph()
            run(p, "What This Involves", size=BODY_SIZE, bold=True, color=BRAND_NAVY)
            para(doc, story["what_involves"])

            # Why it takes the hours it does
            p = doc.add_paragraph()
            run(p, "Why It Takes the Hours It Does", size=BODY_SIZE, bold=True, color=BRAND_NAVY)
            para(doc, story["why_hours"])

            # Cost drivers
            if story.get("cost_drivers"):
                p = doc.add_paragraph()
                run(p, "Cost Drivers", size=BODY_SIZE, bold=True, color=BRAND_NAVY)
                for driver in story["cost_drivers"]:
                    bullet(doc, driver)

            # Why the range exists
            p = doc.add_paragraph()
            run(p, "Why the Range Exists", size=BODY_SIZE, bold=True, color=BRAND_NAVY)
            para(doc, story["why_range"])

            # Confidence and dependencies
            p = doc.add_paragraph()
            run(p, "Confidence: ", size=BODY_SIZE, bold=True, color=BRAND_NAVY)
            conf = story.get("confidence", "Medium")
            conf_color = "2E7D32" if conf == "High" else ("E65100" if conf == "Medium" else "C62828")
            run(p, conf, size=BODY_SIZE, bold=True, color=conf_color)
            if story.get("confidence_reason"):
                run(p, f" — {story['confidence_reason']}", size=BODY_SIZE)

            if story.get("dependencies"):
                p = doc.add_paragraph()
                run(p, "Dependencies: ", size=BODY_SIZE, bold=True, color=BRAND_NAVY)
                run(p, story["dependencies"], size=BODY_SIZE)

            # Separator
            doc.add_paragraph()

        doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # PHASE BREAKDOWN
    # ════════════════════════════════════════════════════════════════════
    heading(doc, "Phase Breakdown Summary")

    para(doc, "Aggregate effort across all stories, broken down by development phase. This view shows where the effort actually lives and where agentic development delivers the most acceleration.")
    doc.add_paragraph()

    phase_data = content.get("phase_breakdown", [])
    phase_rows = [[p["phase"], p["pct"], p["trad_hrs"], p["agnt_hrs"], p["savings"]] for p in phase_data if not p.get("_total")]
    total = [p for p in phase_data if p.get("_total")]
    if total:
        phase_rows.append({"_total": True, "_vals": [total[0]["phase"], total[0]["pct"], total[0]["trad_hrs"], total[0]["agnt_hrs"], total[0]["savings"]]})

    slim_table(doc,
        ["Phase", "% of Effort", "Traditional Hours", "Agentic Hours", "Savings"],
        phase_rows,
        col_widths=[1.8, 1.0, 1.4, 1.2, 0.8],
        right_align_cols=[1, 2, 3, 4]
    )

    if content.get("phase_narrative"):
        doc.add_paragraph()
        para(doc, content["phase_narrative"], italic=True)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # RISK ASSESSMENT
    # ════════════════════════════════════════════════════════════════════
    heading(doc, "Risk Assessment")

    risks = content.get("risks", [])
    if risks:
        slim_table(doc,
            ["Risk", "Level", "Affected Stories", "Hour Impact", "Mitigation"],
            [[r["risk"], r["level"], r["stories"], r["impact"], r["mitigation"]] for r in risks],
            col_widths=[2.2, 0.6, 1.0, 0.8, 2.4]
        )

    doc.add_paragraph()

    risk_buffer = content.get("risk_buffer", {})
    if risk_buffer:
        callout(doc, f"Risk Buffer Applied: +{risk_buffer.get('pct', '17')}% = +{risk_buffer.get('hours', '24')} hours. "
                     f"Reason: {risk_buffer.get('reason', 'Medium confidence overall.')}")

    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════
    # ASSUMPTIONS & EXCLUSIONS
    # ════════════════════════════════════════════════════════════════════
    heading(doc, "Assumptions & Exclusions")

    heading(doc, "Assumptions", level=2)
    para(doc, "This estimate relies on the following conditions. If any assumption is incorrect, the affected stories should be re-estimated.", italic=True)
    doc.add_paragraph()

    assumptions = content.get("assumptions", [])
    if assumptions:
        slim_table(doc,
            ["ID", "Assumption", "If Wrong", "Impact"],
            [[a["id"], a["assumption"], a["if_wrong"], a["impact"]] for a in assumptions],
            col_widths=[0.4, 3.0, 2.0, 1.5]
        )

    doc.add_paragraph()

    heading(doc, "Exclusions (Out of Scope)", level=2)
    for excl in content.get("exclusions", []):
        bullet(doc, excl)

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # HISTORICAL CALIBRATION
    # ════════════════════════════════════════════════════════════════════
    if content.get("calibration"):
        heading(doc, "Historical Calibration")
        for p_text in content["calibration"]:
            para(doc, p_text)
        doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════
    # UNRESOLVED QUESTIONS
    # ════════════════════════════════════════════════════════════════════
    if content.get("unresolved_questions"):
        heading(doc, "Unresolved Questions Affecting Scope")

        callout(doc, f"These {len(content['unresolved_questions'])} questions must be resolved before implementation begins. Each has a quantified impact on the estimate.")
        doc.add_paragraph()

        slim_table(doc,
            ["#", "Question", "Needed From", "Hour Impact If Unresolved", "Blocking?"],
            content["unresolved_questions"],
            col_widths=[0.3, 2.5, 1.2, 1.8, 0.7]
        )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # SUMMARY TABLE
    # ════════════════════════════════════════════════════════════════════
    heading(doc, "Estimate Summary")

    para(doc, "All stories with hour estimates, organized by feature. Hours shown are agentic (AI-assisted) delivery model.")
    doc.add_paragraph()

    summary_rows = content.get("summary_table", [])
    slim_table(doc,
        ["#", "Story", "Feature", "Optimistic", "Most Likely", "Pessimistic", "PERT Expected", "Confidence"],
        [r if not isinstance(r, dict) else {"_total": True, "_vals": r["_vals"]} for r in summary_rows],
        col_widths=[0.3, 2.2, 1.3, 0.7, 0.8, 0.7, 0.8, 0.7],
        right_align_cols=[3, 4, 5, 6]
    )

    doc.add_paragraph()

    # Feature subtotals
    if content.get("feature_totals"):
        heading(doc, "By Feature", level=2)
        slim_table(doc,
            ["Feature", "Stories", "Agentic PERT", "Traditional PERT", "Savings"],
            [r if not isinstance(r, dict) else {"_total": True, "_vals": r["_vals"]} for r in content["feature_totals"]],
            col_widths=[2.5, 0.8, 1.2, 1.3, 0.8],
            right_align_cols=[1, 2, 3, 4]
        )

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════
    # CONFIDENCE SUMMARY
    # ════════════════════════════════════════════════════════════════════
    heading(doc, "Confidence Summary")

    conf_rows = content.get("confidence_summary", [])
    slim_table(doc,
        ["Metric", "Value"],
        conf_rows,
        col_widths=[3.5, 3.5]
    )

    doc.add_paragraph()
    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════
    # SIGN-OFF
    # ════════════════════════════════════════════════════════════════════
    heading(doc, "Approval & Sign-Off")

    para(doc, "By signing below, the undersigned acknowledges that they have reviewed this estimate and agree that it accurately represents the expected scope and effort for the described work.")
    doc.add_paragraph()

    signoff = doc.add_table(rows=4, cols=4)
    signoff.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ci, h in enumerate(["Role", "Name", "Signature", "Date"]):
        cell = signoff.rows[0].cells[ci]
        cell.text = ""
        p = cell.paragraphs[0]
        run(p, h, size=BODY_SIZE, bold=True, color=BRAND_NAVY)
        set_border_bottom(cell, BRAND_PRIMARY)
    for ri, role in enumerate(["Product Owner", "Project Manager", "Technical Lead"]):
        cell = signoff.rows[ri + 1].cells[0]
        cell.text = ""
        p = cell.paragraphs[0]
        run(p, role, size=BODY_SIZE)
        for ci in range(1, 4):
            cell = signoff.rows[ri + 1].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            run(p, "________________________", size=BODY_SIZE, color="CCCCCC")

    # ════════════════════════════════════════════════════════════════════
    # FOOTER
    # ════════════════════════════════════════════════════════════════════
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.clear()
        r = p.add_run(f"{content['project_name']} \u2022 Emergent Software \u2022 Confidential")
        r.font.size = Pt(9)
        r.font.color.rgb = rgb("999999")
        r.font.name = FONT

        p2 = footer.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run()
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        r2._element.append(fldChar1)
        r3 = p2.add_run()
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        r3._element.append(instrText)
        r4 = p2.add_run()
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        r4._element.append(fldChar2)
        for r in [r2, r3, r4]:
            r.font.size = Pt(9)
            r.font.color.rgb = rgb("999999")

    # Save
    output_path = content["output_path"]
    doc.save(output_path)
    print(f"Estimate document saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python generate_estimate_doc.py estimate_content.json")
        sys.exit(1)
    with open(sys.argv[1], encoding="utf-8") as f:
        data = json.load(f)
    generate_estimate(data)
