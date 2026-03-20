"""
Generate a branded status report .docx from JSON content.

Usage:
    python generate_status_report.py content.json

The JSON file should have this structure:
{
    "date_range": "March 4 – March 18, 2026",
    "output_path": "/path/to/ProjectName-StatusUpdate-03182026.docx",
    "branding": {
        "project_name": "Honda AIM",
        "primary_color": "2E7D32",
        "accent_color": "1565C0",
        "light_color": "E8F5E9",
        "footer_text": "Honda AIM • Emergent Software • Confidential"
    },
    "executive_summary": "Summary paragraph text...",
    "sections": [
        {
            "heading": "Section Name",
            "items": ["Bullet point 1", "Bullet point 2"]
        }
    ],
    "callout": {
        "text": "COMING UP — description of upcoming focus area"
    },
    "timeline": [
        {"phase": "Phase 1", "focus": "...", "dates": "Completed", "outcome": "..."}
    ],
    "blockers": ["Blocker 1", "Blocker 2"],
    "next_steps": ["Step 1", "Step 2"],
    "pull_requests": {
        "date_range_label": "3/4 – 3/18",
        "items": ["PR #269: Description", "PR #270: Description"]
    }
}

Branding defaults (used when "branding" is absent or a field is missing):
    primary_color:  2E7D32  (Emergent green)
    accent_color:   1565C0  (Emergent blue)
    light_color:    E8F5E9  (light green)
    project_name:   "Status Report"
    footer_text:    "Status Report • Emergent Software • Confidential"
"""

import json
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# Emergent fallback brand colors (extracted from EmergentStyles.docx)
_EMERGENT_PRIMARY = "EE3342"   # Emergent coral/red — headings, header bar
_EMERGENT_ACCENT = "333F4F"    # Dark navy — accent bar, title text
_EMERGENT_LIGHT = "FDF4F4"     # Very light coral tint — alternating table rows

# Fixed colors independent of project branding
CALLOUT_YELLOW = "FFF2CC"
TEXT_COLOR = "333333"
WHITE = "FFFFFF"
LIGHT_DATE = "F2D4D7"

# Font settings (Emergent uses Calibri Light)
BODY_FONT = "Calibri Light"
BODY_SIZE = 11
BULLET_SIZE = 11


def _resolve_brand(content: dict) -> dict:
    """
    Extract brand colors and labels from the JSON content's 'branding' object.
    Falls back to Emergent defaults for any missing field.
    """
    b = content.get("branding") or {}
    project_name = b.get("project_name") or "Status Report"
    return {
        "project_name": project_name,
        "primary": b.get("primary_color") or _EMERGENT_PRIMARY,
        "accent": b.get("accent_color") or _EMERGENT_ACCENT,
        "light": b.get("light_color") or _EMERGENT_LIGHT,
        "footer_text": b.get("footer_text") or f"{project_name} \u2022 Emergent Software \u2022 Confidential",
    }


def set_cell_shading(cell, color: str) -> None:
    """Set the background shading of a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._element.get_or_add_tcPr().append(shading)


def set_cell_margins(cell, top: int = 0, bottom: int = 0, start: int = 0, end: int = 0) -> None:
    """Set cell margins in DXA units (twentieths of a point)."""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:start w:w="{start}" w:type="dxa"/>'
        f'  <w:end w:w="{end}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)


def set_cell_border_bottom(cell, color: str, sz: int = 8) -> None:
    """Set a bottom border on a table cell."""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)


def add_run(paragraph, text: str, font_name: str = None, font_size: float = None, bold: bool = None, color: str = None):
    """Add a formatted run to a paragraph."""
    run = paragraph.add_run(text)
    if font_name:
        run.font.name = font_name
    if font_size:
        run.font.size = Pt(font_size)
    if bold is not None:
        run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def _remove_table_borders(table) -> None:
    """Remove all borders from a table element."""
    tbl = table._element
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


def _set_table_full_width(table) -> None:
    """Stretch a table to 100% of the page width."""
    tbl = table._element
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
    tblPr.append(tblW)


def create_header(doc, date_range: str, brand: dict):
    """Create a clean header — title on left, logo on right, thin accent line."""
    # Use a 2-column table: left = title/subtitle/date, right = logo
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_full_width(table)
    _remove_table_borders(table)

    # Left cell: title, subtitle, date
    left = table.cell(0, 0)
    set_cell_margins(left, top=0, bottom=0, start=0, end=0)

    p_title = left.paragraphs[0]
    p_title.paragraph_format.space_after = Pt(2)
    add_run(p_title, brand["project_name"], font_name=BODY_FONT, font_size=28, bold=True, color=brand["accent"])

    p_sub = left.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(2)
    subtitle = brand.get("report_subtitle", "Status Update")
    add_run(p_sub, subtitle, font_name=BODY_FONT, font_size=14, color=brand["primary"])

    p_date = left.add_paragraph()
    p_date.paragraph_format.space_before = Pt(0)
    p_date.paragraph_format.space_after = Pt(0)
    add_run(p_date, date_range, font_name=BODY_FONT, font_size=11, color="666666")

    # Right cell: logo, right-aligned
    right = table.cell(0, 1)
    set_cell_margins(right, top=0, bottom=0, start=0, end=0)

    # Set right cell width narrower
    tc = right._element
    tcPr = tc.get_or_add_tcPr()
    tcW = parse_xml(f'<w:tcW {nsdecls("w")} w:w="2500" w:type="dxa"/>')
    tcPr.append(tcW)

    logo_path = Path(__file__).parent / "emergent-logo.png"
    p_logo = right.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_logo.paragraph_format.space_before = Pt(4)
    if logo_path.exists():
        p_logo.add_run().add_picture(str(logo_path), width=Inches(1.8))

    # Thin accent line below the header table
    line_table = doc.add_table(rows=1, cols=1)
    _set_table_full_width(line_table)
    line_cell = line_table.cell(0, 0)
    run = line_cell.paragraphs[0].add_run("")
    run.font.size = Pt(1)
    tbl = line_table._element
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:bottom w:val="single" w:sz="12" w:space="0" w:color="{brand["primary"]}"/>'
        '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)


def create_callout_table(doc, text: str):
    """Create a yellow callout box."""
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_YELLOW)
    set_cell_margins(cell, top=100, bottom=100, start=150, end=150)
    p = cell.paragraphs[0]

    # Bold the prefix up to the first colon or em-dash
    split_char = None
    if ":" in text:
        split_char = ":"
    elif "\u2014" in text:
        split_char = "\u2014"

    if split_char:
        prefix, rest = text.split(split_char, 1)
        add_run(p, prefix + split_char, bold=True, font_size=BODY_SIZE, color=TEXT_COLOR)
        add_run(p, rest, font_size=BODY_SIZE, color=TEXT_COLOR)
    else:
        add_run(p, text, font_size=BODY_SIZE, color=TEXT_COLOR)

    _set_table_full_width(table)
    return table


def create_timeline_table(doc, phases: list, brand: dict):
    """Create the schedule/timeline table with branded header row."""
    cols = ["Phase", "Focus", "Target Dates", "Outcome"]
    table = doc.add_table(rows=1 + len(phases), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_full_width(table)

    # Header row
    for ci, col_name in enumerate(cols):
        cell = table.cell(0, ci)
        set_cell_shading(cell, brand["primary"])
        p = cell.paragraphs[0]
        add_run(p, col_name, bold=True, font_size=11, color=WHITE)

    # Data rows with alternating shading
    for ri, phase in enumerate(phases):
        values = [
            phase.get("phase", ""),
            phase.get("focus", ""),
            phase.get("dates", ""),
            phase.get("outcome", ""),
        ]
        for ci, val in enumerate(values):
            cell = table.cell(ri + 1, ci)
            if ri % 2 == 1:
                set_cell_shading(cell, brand["light"])
            add_run(cell.paragraphs[0], val, font_size=11, color=TEXT_COLOR)

    return table


def create_footer_table(doc, brand: dict):
    """Create the branded footer bar."""
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, brand["primary"])
    set_cell_margins(cell, top=60, bottom=60, start=150, end=150)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, brand["footer_text"], font_size=8, color=WHITE)
    _set_table_full_width(table)
    return table


def _set_header_row_repeat(table):
    """Set the first row of a table to repeat as header across pages."""
    row = table.rows[0]
    trPr = row._element.get_or_add_trPr()
    trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))


def _style_header_row(table, cols: list, brand: dict, right_align_from: int = 1):
    """Style a header row with thin bottom border and colored text — no solid fill."""
    for ci, col_name in enumerate(cols):
        cell = table.cell(0, ci)
        set_cell_margins(cell, top=40, bottom=40, start=80, end=80)
        # Thin bottom border in primary color
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        tcBorders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="8" w:space="0" w:color="{brand["primary"]}"/>'
            '</w:tcBorders>'
        )
        tcPr.append(tcBorders)
        p = cell.paragraphs[0]
        if ci >= right_align_from:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_run(p, col_name, bold=True, font_size=11, color=brand["primary"], font_name=BODY_FONT)
    _set_header_row_repeat(table)


def _style_totals_row(table, row_idx: int, values: list, brand: dict, right_align_from: int = 1):
    """Style a totals row with top border and bold primary-colored text — no solid fill."""
    for ci, val in enumerate(values):
        cell = table.cell(row_idx, ci)
        set_cell_margins(cell, top=40, bottom=40, start=80, end=80)
        # Thin top border
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        tcBorders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="8" w:space="0" w:color="{brand["primary"]}"/>'
            '</w:tcBorders>'
        )
        tcPr.append(tcBorders)
        p = cell.paragraphs[0]
        if ci >= right_align_from:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        add_run(p, str(val), font_size=11, color=brand["primary"], font_name=BODY_FONT, bold=True)


def create_epic_summary_table(doc, epics: list, brand: dict):
    """Create a table showing Epic-level completion with % progress and totals row."""
    cols = ["Epic", "Done", "Total", "% Complete"]
    table = doc.add_table(rows=2 + len(epics), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_full_width(table)
    _remove_table_borders(table)

    _style_header_row(table, cols, brand)

    # Data rows with alternating shading
    sum_done = 0
    sum_total = 0
    for ri, epic in enumerate(epics):
        done = epic.get("done", 0)
        total = epic.get("total", 0)
        pct = epic.get("pct", 0)
        sum_done += done
        sum_total += total
        pct_str = f"{pct}%"
        values = [epic.get("epic", ""), str(done), str(total), pct_str]
        for ci, val in enumerate(values):
            cell = table.cell(ri + 1, ci)
            set_cell_margins(cell, top=40, bottom=40, start=80, end=80)
            if ri % 2 == 1:
                set_cell_shading(cell, brand["light"])
            p = cell.paragraphs[0]
            if ci >= 1:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if ci == 3:
                clr = brand["primary"] if pct == 100 else TEXT_COLOR
                add_run(p, val, font_size=11, color=clr, font_name=BODY_FONT, bold=(pct == 100))
            else:
                add_run(p, val, font_size=11, color=TEXT_COLOR, font_name=BODY_FONT)

    # Totals row
    total_row = len(epics) + 1
    overall_pct = (sum_done * 100 // sum_total) if sum_total > 0 else 0
    _style_totals_row(table, total_row, ["Project Total", str(sum_done), str(sum_total), f"{overall_pct}%"], brand)

    return table


def create_velocity_table(doc, velocity: dict, brand: dict):
    """Create a velocity metrics table showing this period vs overall with trend."""
    cols = ["Metric", "This Period", "Project Avg", "Trend"]
    rows_data = velocity.get("rows", [])
    has_totals = bool(velocity.get("totals"))
    table = doc.add_table(rows=1 + len(rows_data) + (1 if has_totals else 0), cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_full_width(table)
    _remove_table_borders(table)

    _style_header_row(table, cols, brand, right_align_from=1)

    # Data rows
    for ri, row in enumerate(rows_data):
        trend = row.get("trend", "")
        values = [row.get("metric", ""), row.get("period", ""), row.get("overall", ""), trend]
        for ci, val in enumerate(values):
            cell = table.cell(ri + 1, ci)
            set_cell_margins(cell, top=40, bottom=40, start=80, end=80)
            if ri % 2 == 1:
                set_cell_shading(cell, brand["light"])
            p = cell.paragraphs[0]
            if ci >= 1:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if ci == 3 and val.startswith("\u25b2"):
                add_run(p, val, font_size=11, color="2E7D32", font_name=BODY_FONT)
            elif ci == 3 and val.startswith("\u25bc"):
                add_run(p, val, font_size=11, color=brand["primary"], font_name=BODY_FONT)
            else:
                add_run(p, str(val), font_size=11, color=TEXT_COLOR, font_name=BODY_FONT)

    # Totals row
    if velocity.get("totals"):
        t = velocity["totals"]
        total_row = len(rows_data) + 1
        _style_totals_row(table, total_row,
                          [t.get("metric", "Total"), t.get("period", ""), t.get("overall", ""), t.get("trend", "")],
                          brand, right_align_from=1)

    return table


def _velocity_trend(period_vel: float, project_vel: float) -> str:
    """Calculate a velocity trend indicator comparing period to project average."""
    if project_vel == 0:
        return "\u25b2 New" if period_vel > 0 else "\u2014"
    pct_change = ((period_vel - project_vel) / project_vel) * 100
    if pct_change > 20:
        return f"\u25b2 +{pct_change:.0f}%"
    elif pct_change > 0:
        return f"\u25b2 +{pct_change:.0f}%"
    elif pct_change > -20:
        return f"\u25bc {pct_change:.0f}%"
    else:
        return f"\u25bc {pct_change:.0f}%"


def create_developer_stats_table(doc, devs: list, totals: dict, brand: dict):
    """Create a per-developer stats table with velocity and trend. Stories first."""
    INACTIVE_COLOR = "AAAAAA"
    cols = ["Developer", "Stories", "Proj Stories", "Points", "PRs", "Trend"]
    table = doc.add_table(rows=2 + len(devs), cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_full_width(table)
    _remove_table_borders(table)

    proj_weeks = totals.get("project_weeks", 1) or 1
    period_weeks = totals.get("period_weeks", 1) or 1

    _style_header_row(table, cols, brand)

    # Sort: active devs by period stories desc, then inactive at bottom
    active = [d for d in devs if d.get("period_stories", 0) > 0 or d.get("period_pts", 0) > 0 or d.get("period_prs", 0) > 0]
    inactive = [d for d in devs if d not in active]
    active.sort(key=lambda d: (d.get("period_stories", 0), d.get("period_pts", 0)), reverse=True)
    sorted_devs = active + inactive

    # Developer rows
    for ri, dev in enumerate(sorted_devs):
        period_pts = dev.get("period_pts", 0)
        project_pts = dev.get("project_pts", 0)
        period_stories = dev.get("period_stories", 0)
        project_stories = dev.get("project_stories", 0)
        period_prs = dev.get("period_prs", 0)
        project_prs = dev.get("project_prs", 0)
        period_story_vel = round(period_stories / period_weeks, 1)
        project_story_vel = round(project_stories / proj_weeks, 1) if project_stories else 0
        is_inactive = dev in inactive
        trend = _velocity_trend(period_story_vel, project_story_vel) if not is_inactive else "\u2014"
        pts_combined = str(int(period_pts)) + "/" + str(int(project_pts))
        prs_combined = str(period_prs) + "/" + str(project_prs)
        values = [
            dev["name"],
            str(period_stories),
            str(project_stories),
            pts_combined,
            prs_combined,
            trend,
        ]
        text_color = INACTIVE_COLOR if is_inactive else TEXT_COLOR
        for ci, val in enumerate(values):
            cell = table.cell(ri + 1, ci)
            set_cell_margins(cell, top=40, bottom=40, start=80, end=80)
            if is_inactive:
                set_cell_shading(cell, "F0F0F0")
            elif ri % 2 == 1:
                set_cell_shading(cell, brand["light"])
            p = cell.paragraphs[0]
            if ci >= 1:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            # Color trend indicators
            if ci == 5 and not is_inactive and val.startswith("\u25b2"):
                add_run(p, val, font_size=11, color="2E7D32", font_name=BODY_FONT)
            elif ci == 5 and not is_inactive and val.startswith("\u25bc"):
                add_run(p, val, font_size=11, color=brand["primary"], font_name=BODY_FONT)
            else:
                add_run(p, val, font_size=11, color=text_color, font_name=BODY_FONT)

    # Totals row
    total_row = len(sorted_devs) + 1
    period_total = totals.get("period_pts", 0)
    project_total = totals.get("project_pts", 0)
    total_period_stories = totals.get("period_stories", sum(d.get("period_stories", 0) for d in devs))
    total_project_stories = totals.get("project_stories", sum(d.get("project_stories", 0) for d in devs))
    total_period_prs = totals.get("period_prs", 0)
    total_project_prs = totals.get("project_prs", 0)
    total_story_vel_period = round(total_period_stories / period_weeks, 1)
    total_story_vel_project = round(total_project_stories / proj_weeks, 1) if total_project_stories else 0
    total_trend = _velocity_trend(total_story_vel_period, total_story_vel_project)
    _style_totals_row(table, total_row, [
        "Team Total",
        str(total_period_stories),
        str(total_project_stories),
        str(int(period_total)) + "/" + str(int(project_total)),
        str(total_period_prs) + "/" + str(total_project_prs),
        total_trend,
    ], brand)

    return table


def _configure_styles(doc, brand: dict) -> None:
    """Apply default font and heading styles to the document."""
    # Normal style
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(BODY_SIZE)
    style.font.color.rgb = RGBColor.from_string(TEXT_COLOR)

    # Heading styles — sized for a clean, compact look
    for level, size, sp_before, sp_after in [(1, 16, 12, 4), (2, 13, 8, 4)]:
        h_style = doc.styles[f"Heading {level}"]
        h_style.font.color.rgb = RGBColor.from_string(brand["primary"])
        h_style.font.bold = True
        h_style.font.size = Pt(size)
        h_style.font.name = BODY_FONT
        h_style.paragraph_format.space_before = Pt(sp_before)
        h_style.paragraph_format.space_after = Pt(sp_after)

    # List Bullet style: breathing room, hanging indent, at-least line spacing
    bullet_style = doc.styles["List Bullet"]
    bullet_style.font.name = BODY_FONT
    bullet_style.font.size = Pt(BULLET_SIZE)
    bullet_style.paragraph_format.space_before = Pt(6)
    bullet_style.paragraph_format.space_after = Pt(6)
    bullet_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
    bullet_style.paragraph_format.line_spacing = Pt(14)
    bullet_style.paragraph_format.left_indent = Inches(0.5)
    bullet_style.paragraph_format.first_line_indent = Inches(-0.25)

    # Disable "don't add space between same-style paragraphs" for bullet lists
    pPr = bullet_style.element.get_or_add_pPr()
    ctx_spacing = pPr.find(qn("w:contextualSpacing"))
    if ctx_spacing is not None:
        pPr.remove(ctx_spacing)
    pPr.append(parse_xml(f'<w:contextualSpacing {nsdecls("w")} w:val="0"/>'))


def generate_report(content: dict) -> str:
    """Generate the status report document and return the output path."""
    brand = _resolve_brand(content)

    # Always start from a blank document — this ensures all built-in styles
    # (Normal, List Bullet, Heading 1, etc.) are available
    doc = Document()

    _configure_styles(doc, brand)

    # Page margins and footer
    for section in doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)

        # Page footer — thin line + address + website
        footer = section.footer
        footer.is_linked_to_previous = False

        # Thin line above footer text
        p_line = footer.paragraphs[0]
        p_line.paragraph_format.space_after = Pt(4)
        pPr = p_line._element.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="4" w:space="4" w:color="{brand["primary"]}"/>'
            '</w:pBdr>'
        )
        pPr.append(pBdr)

        footer_text = brand.get("footer_text", "")
        if not footer_text:
            footer_text = f'{brand["project_name"]} \u2022 Emergent Software \u2022 Confidential'
        p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p_line, "2038 Ford Pkwy, Suite 439, Saint Paul, MN 55116  |  emergentsoftware.net",
                font_name=BODY_FONT, font_size=11, color="999999")

        # Page number line
        p_page = footer.add_paragraph()
        p_page.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_page.paragraph_format.space_before = Pt(2)
        run_pre = add_run(p_page, "Page ", font_name=BODY_FONT, font_size=11, color="999999")
        # Insert PAGE field
        fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run_field = p_page.add_run()
        run_field.font.name = BODY_FONT
        run_field.font.size = Pt(8)
        run_field.font.color.rgb = RGBColor.from_string("999999")
        run_field._element.append(fldChar1)
        instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run_field2 = p_page.add_run()
        run_field2._element.append(instrText)
        fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
        run_field3 = p_page.add_run()
        run_field3._element.append(fldChar2)
        run_field4 = p_page.add_run("1")
        run_field4.font.size = Pt(8)
        run_field4.font.color.rgb = RGBColor.from_string("999999")
        fldChar3 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run_field5 = p_page.add_run()
        run_field5._element.append(fldChar3)
        add_run(p_page, " of ", font_name=BODY_FONT, font_size=11, color="999999")
        # Insert NUMPAGES field
        fldChar4 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run_np1 = p_page.add_run()
        run_np1._element.append(fldChar4)
        instrText2 = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> NUMPAGES </w:instrText>')
        run_np2 = p_page.add_run()
        run_np2._element.append(instrText2)
        fldChar5 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
        run_np3 = p_page.add_run()
        run_np3._element.append(fldChar5)
        run_np4 = p_page.add_run("1")
        run_np4.font.size = Pt(8)
        run_np4.font.color.rgb = RGBColor.from_string("999999")
        fldChar6 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run_np5 = p_page.add_run()
        run_np5._element.append(fldChar6)

    # Header
    create_header(doc, content["date_range"], brand)

    # Executive Summary
    doc.add_heading("Executive Summary", level=1)
    p = doc.add_paragraph(content["executive_summary"])
    p.style = doc.styles["Normal"]

    # Key Achievements
    doc.add_heading("Key Achievements", level=1)
    for section in content.get("sections", []):
        doc.add_heading(section["heading"], level=2)
        for i, item in enumerate(section.get("items", [])):
            if i == 0:
                # First item is a summary paragraph, not a bullet
                doc.add_paragraph(item)
            else:
                doc.add_paragraph(item, style="List Bullet")

    # Callout box
    if content.get("callout"):
        doc.add_paragraph()
        create_callout_table(doc, content["callout"]["text"])
        doc.add_paragraph()

    # Epic Progress table
    if content.get("epic_summary"):
        doc.add_heading("Epic Progress", level=1)
        create_epic_summary_table(doc, content["epic_summary"], brand)

    # Remaining work discussion
    if content.get("remaining_work"):
        doc.add_heading("Remaining Work", level=2)
        for item in content["remaining_work"]:
            doc.add_paragraph(item, style="List Bullet")
        if content.get("remaining_work_footnote"):
            p = doc.add_paragraph()
            run = p.add_run(content["remaining_work_footnote"])
            run.italic = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # Velocity
    if content.get("velocity"):
        doc.add_heading("Velocity", level=1)
        create_velocity_table(doc, content["velocity"], brand)
        if content["velocity"].get("summary"):
            doc.add_heading("Summary", level=2)
            for item in content["velocity"]["summary"] if isinstance(content["velocity"]["summary"], list) else [content["velocity"]["summary"]]:
                doc.add_paragraph(item, style="List Bullet")

    # Developer Stats
    if content.get("developer_stats"):
        doc.add_heading("Developer Metrics", level=1)
        totals = content.get("developer_totals", {})
        create_developer_stats_table(doc, content["developer_stats"], totals, brand)
        # Table legend — small italic note
        legend = doc.add_paragraph()
        legend.paragraph_format.space_before = Pt(2)
        legend.paragraph_format.space_after = Pt(4)
        note = legend.add_run("Stories = completed (Dev Complete, Resolved, or Closed). Points and PRs shown as period/project totals. Trend compares story velocity (stories/week) against project average.")
        note.italic = True
        note.font.size = Pt(9)
        note.font.color.rgb = RGBColor.from_string("999999")
        note.font.name = BODY_FONT
        if content.get("developer_summary"):
            doc.add_heading("Summary", level=2)
            for item in content["developer_summary"] if isinstance(content["developer_summary"], list) else [content["developer_summary"]]:
                doc.add_paragraph(item, style="List Bullet")

    # Code Quality & Technical Insights
    if content.get("code_quality"):
        cq = content["code_quality"]
        doc.add_heading("Code Quality & Technical Insights", level=1)
        # Summary table
        tbl = doc.add_table(rows=2, cols=5)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ["PRs Reviewed", "Clean Rate", "Critical", "Major", "Minor"]
        values = [
            str(cq.get("prs_reviewed", 0)),
            f"{cq.get('clean_rate', 0)}%",
            str(cq.get("critical", 0)),
            str(cq.get("major", 0)),
            str(cq.get("minor", 0))
        ]
        _remove_table_borders(tbl)
        _style_header_row(tbl, headers, brand, right_align_from=2)
        for i, v in enumerate(values):
            cell = tbl.rows[1].cells[i]
            set_cell_margins(cell, top=40, bottom=40, start=80, end=80)
            p = cell.paragraphs[0]
            if i >= 2:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            add_run(p, v, font_size=11, font_name=BODY_FONT)

        if cq.get("summary"):
            doc.add_heading("Summary", level=2)
            for item in cq["summary"] if isinstance(cq["summary"], list) else [cq["summary"]]:
                doc.add_paragraph(item, style="List Bullet")

        if cq.get("notable"):
            doc.add_heading("Notable Issues Addressed", level=2)
            for item in cq["notable"]:
                doc.add_paragraph(item, style="List Bullet")

    # Timeline / Schedule
    if content.get("timeline"):
        doc.add_heading("Schedule / Timeline", level=1)
        create_timeline_table(doc, content["timeline"], brand)

    # Burndown Analysis
    if content.get("burndown"):
        doc.add_heading("Burndown Analysis", level=1)
        bd = content["burndown"]
        # Show last 8 weeks max to keep it compact
        display_bd = bd[-8:] if len(bd) > 8 else bd
        cols_bd = ["Week Of", "Completed", "Cumulative", "Remaining", "% Done"]
        tbl_bd = doc.add_table(rows=1 + len(display_bd), cols=5)
        tbl_bd.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set_table_full_width(tbl_bd)
        _remove_table_borders(tbl_bd)
        _style_header_row(tbl_bd, cols_bd, brand, right_align_from=1)
        for ri, row_data in enumerate(display_bd):
            vals = [row_data["week"], str(row_data["completed"]), str(row_data["cumulative"]),
                    str(row_data["remaining"]), str(row_data["pct"]) + "%"]
            for ci, val in enumerate(vals):
                cell = tbl_bd.cell(ri + 1, ci)
                set_cell_margins(cell, top=40, bottom=40, start=80, end=80)
                if ri % 2 == 1:
                    set_cell_shading(cell, brand["light"])
                p = cell.paragraphs[0]
                if ci >= 1:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                add_run(p, val, font_size=11, color=TEXT_COLOR, font_name=BODY_FONT)
        if content.get("burndown_summary"):
            p_bs = doc.add_paragraph()
            p_bs.paragraph_format.space_before = Pt(4)
            run_bs = p_bs.add_run(content["burndown_summary"])
            run_bs.font.size = Pt(11)
            run_bs.font.name = BODY_FONT

    # Period Comparison
    if content.get("comparison"):
        comp = content["comparison"]
        doc.add_heading("Period Comparison", level=1)
        cols_cmp = ["Metric", comp["prev"]["label"], comp["curr"]["label"], "Change"]
        tbl_cmp = doc.add_table(rows=3, cols=4)
        tbl_cmp.alignment = WD_TABLE_ALIGNMENT.CENTER
        _set_table_full_width(tbl_cmp)
        _remove_table_borders(tbl_cmp)
        _style_header_row(tbl_cmp, cols_cmp, brand, right_align_from=1)
        comp_rows = [
            ["Stories Completed", str(comp["prev"]["stories"]), str(comp["curr"]["stories"]),
             ("{:+d}".format(comp["delta_stories"]))],
            ["Points Delivered", str(int(comp["prev"]["pts"])), str(int(comp["curr"]["pts"])),
             ("{:+.0f}".format(comp["delta_pts"]))],
        ]
        for ri, row_vals in enumerate(comp_rows):
            for ci, val in enumerate(row_vals):
                cell = tbl_cmp.cell(ri + 1, ci)
                set_cell_margins(cell, top=40, bottom=40, start=80, end=80)
                if ri % 2 == 1:
                    set_cell_shading(cell, brand["light"])
                p = cell.paragraphs[0]
                if ci >= 1:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                # Color the delta
                if ci == 3 and val.startswith("+"):
                    add_run(p, val, font_size=11, color="2E7D32", font_name=BODY_FONT)
                elif ci == 3 and val.startswith("-"):
                    add_run(p, val, font_size=11, color=brand["primary"], font_name=BODY_FONT)
                else:
                    add_run(p, val, font_size=11, color=TEXT_COLOR, font_name=BODY_FONT)
        if content.get("comparison_summary"):
            p_cs = doc.add_paragraph()
            p_cs.paragraph_format.space_before = Pt(4)
            run_cs = p_cs.add_run(content["comparison_summary"])
            run_cs.font.size = Pt(11)
            run_cs.font.name = BODY_FONT

    # QA Pipeline
    if content.get("qa_pipeline"):
        qa = content["qa_pipeline"]
        doc.add_heading("QA Pipeline", level=1)
        p_qa_intro = doc.add_paragraph()
        run_qi = p_qa_intro.add_run(str(len(qa["stories"])) + " stories currently in QA Testing.")
        run_qi.font.size = Pt(11)
        run_qi.font.name = BODY_FONT
        run_qi.bold = True
        if qa["stories"]:
            cols_qa = ["ID", "Title", "Epic", "Developer", "In QA Since"]
            tbl_qa = doc.add_table(rows=1 + len(qa["stories"]), cols=5)
            tbl_qa.alignment = WD_TABLE_ALIGNMENT.CENTER
            _set_table_full_width(tbl_qa)
            _remove_table_borders(tbl_qa)
            _style_header_row(tbl_qa, cols_qa, brand, right_align_from=4)
            for ri, story in enumerate(qa["stories"]):
                vals = ["#" + str(story["id"]), story["title"][:55], story.get("epic", ""),
                        story.get("dev", ""), story.get("since", "")]
                for ci, val in enumerate(vals):
                    cell = tbl_qa.cell(ri + 1, ci)
                    set_cell_margins(cell, top=40, bottom=40, start=80, end=80)
                    if ri % 2 == 1:
                        set_cell_shading(cell, brand["light"])
                    p = cell.paragraphs[0]
                    add_run(p, val, font_size=11, color=TEXT_COLOR, font_name=BODY_FONT)
        if qa.get("summary"):
            p_qs = doc.add_paragraph()
            p_qs.paragraph_format.space_before = Pt(4)
            run_qs = p_qs.add_run(qa["summary"])
            run_qs.font.size = Pt(11)
            run_qs.font.name = BODY_FONT

    # Blockers / Risks
    doc.add_heading("Blockers / Risks", level=1)
    for item in content.get("blockers", []):
        doc.add_paragraph(item, style="List Bullet")

    # Next Steps
    doc.add_heading("Next Steps / Remaining Work", level=1)
    for item in content.get("next_steps", []):
        doc.add_paragraph(item, style="List Bullet")

    # Appendix A: Completed Stories
    if content.get("appendix_stories"):
        doc.add_heading("Appendix A \u2014 Completed Stories", level=1)
        stories = content["appendix_stories"]
        cols = ["ID", "Title", "Epic", "Developer", "Pts", "State"]
        tbl = doc.add_table(rows=1 + len(stories), cols=len(cols))
        tbl.style = "Table Grid"
        tbl.alignment = 1  # center
        # Header
        for ci, col_name in enumerate(cols):
            cell = tbl.rows[0].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(col_name)
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = "Calibri Light"
            if brand.get("primary"):
                run.font.color.rgb = RGBColor.from_string(brand["primary"])
            p.alignment = 2 if ci == 4 else 0  # right-align Pts
            set_cell_border_bottom(cell, brand.get("primary", "333F4F"))
        # Set header row to repeat across pages
        tbl.rows[0]._tr.get_or_add_trPr().append(
            parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))
        # Rows
        for ri, story in enumerate(stories):
            row = tbl.rows[ri + 1]
            vals = [f"#{story['id']}", story['title'][:60], story['epic'].split(':')[0] if ':' in story['epic'] else story['epic'][:20], story['developer'], str(story['points']), story['state']]
            for ci, val in enumerate(vals):
                cell = row.cells[ci]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(val)
                run.font.size = Pt(11)
                run.font.name = "Calibri Light"
                p.alignment = 2 if ci == 4 else 0
            if ri % 2 == 0:
                for cell in row.cells:
                    set_cell_shading(cell, brand.get("light", "F5F5F5"))
        # Set column widths
        for row in tbl.rows:
            row.cells[0].width = Inches(0.5)
            row.cells[1].width = Inches(2.8)
            row.cells[2].width = Inches(1.2)
            row.cells[3].width = Inches(1.2)
            row.cells[4].width = Inches(0.4)
            row.cells[5].width = Inches(0.9)
        tbl.rows[0].height = Inches(0.25)

    # Appendix B: Pull Requests
    if content.get("pull_requests"):
        pr_data = content["pull_requests"]
        appendix_label = "Appendix B" if content.get("appendix_stories") else "Appendix A"
        doc.add_heading(
            f'{appendix_label} \u2014 Pull Requests Merged ({pr_data["date_range_label"]})',
            level=1,
        )
        for item in pr_data.get("items", []):
            doc.add_paragraph(item, style="List Bullet")

    output_path = content["output_path"]
    doc.save(output_path)
    print(f"Report saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python generate_status_report.py content.json")
        sys.exit(1)

    content_file = sys.argv[1]
    with open(content_file, encoding="utf-8") as f:
        loaded_content = json.load(f)

    generate_report(loaded_content)
